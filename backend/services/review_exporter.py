"""
Review-package exporter.

Produces a zip bundle for business reviewers containing:
  1. 审查报告.docx               — narrative Word report with embedded SVG diagrams
  2. 审查意见表.xlsx             — structured Excel workbook (5 sheets) for feedback
  3. 审查图集.drawio             — multi-page editable diagrams
  4. 图片快照/*.svg              — static SVG per diagram page (embedded in Word too)
  5. 附件_源数据.json            — raw JSON source data for technical reviewers
  6. 使用说明.md                 — usage guide

Input: resolved M1 package and M2 package dicts + m1→m2 mapping list.
Output: (zip_bytes, suggested_filename)
"""
from __future__ import annotations

import io
import json
import zipfile
from datetime import datetime
from typing import Optional

from backend.services import diagram_exporter


# ==============================================================
# Public entry
# ==============================================================

def export_review_package(
    m1_model: dict,
    m2_model: Optional[dict],
) -> tuple[bytes, str]:
    """Build the review package and return (zip_bytes, filename).

    m1_model / m2_model are model dicts (from store.get_model(...).model_dump()).
    m2_model may be None if no M2 has been derived — we still export the M1 part.
    """
    m1_pkg = _latest_package(m1_model)
    m2_pkg = _latest_package(m2_model) if m2_model else None
    m1_class_mappings = _extract_m1_mappings(m1_pkg, m2_pkg)

    date_stamp = datetime.now().strftime("%Y%m%d")
    m1_label = (m1_model.get("label") or m1_model.get("name") or "M1").strip()
    # sanitize for filename
    safe_label = "".join(c for c in m1_label if c.isalnum() or c in "-_·").strip() or "model"
    zip_name = f"审查包_{safe_label}_{date_stamp}.zip"

    # --- Build each artifact ---
    excel_bytes = _build_excel(m1_model, m2_model, m1_pkg, m2_pkg, m1_class_mappings)

    diagrams = {"drawio_xml": "", "svg_by_name": {}, "pdf_bytes": None}
    if m2_pkg:
        diagrams = diagram_exporter.build_review_diagrams(m1_pkg, m2_pkg, m1_class_mappings)

    word_bytes = _build_word(
        m1_model, m2_model, m1_pkg, m2_pkg, m1_class_mappings,
        svg_by_name=diagrams.get("svg_by_name", {}),
    )

    source_json = _build_source_json(m1_model, m2_model, m1_class_mappings)
    readme = _build_readme(m1_model, m2_model)

    # --- Pack into zip ---
    # NOTE: text files intended to be opened in Windows Notepad (.md / .txt) need
    # a UTF-8 BOM (EF BB BF) prefix. Without BOM, Notepad falls back to the system
    # ANSI codepage (GBK on Chinese Windows) and renders UTF-8 as mojibake.
    # Word/Excel internals are UTF-8 already (inside the zip container), no BOM needed.
    # SVG/drawio/JSON have explicit encoding signals (<?xml ... encoding="UTF-8"?> or
    # universal JSON-UTF-8 convention), so no BOM there either (BOM in JSON is illegal
    # per RFC 8259 and breaks strict parsers).
    UTF8_BOM = "\ufeff"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("审查报告.docx", word_bytes)
        zf.writestr("审查意见表.xlsx", excel_bytes)
        if diagrams.get("drawio_xml"):
            zf.writestr("审查图集.drawio", diagrams["drawio_xml"])
        if diagrams.get("pdf_bytes"):
            # PDF version — works without draw.io installed, any PDF viewer can open it
            zf.writestr("审查图集.pdf", diagrams["pdf_bytes"])
        for fname, svg_str in diagrams.get("svg_by_name", {}).items():
            zf.writestr(f"图片快照/{fname}", svg_str.encode("utf-8"))
        zf.writestr("附件_源数据.json", source_json)
        # BOM for the readme so Notepad recognizes UTF-8
        zf.writestr("使用说明.md", UTF8_BOM + readme)

    return buf.getvalue(), zip_name


# ==============================================================
# Helpers to resolve model data
# ==============================================================

def _latest_package(model: Optional[dict]) -> dict:
    if not model:
        return {}
    versions = model.get("versions", []) or []
    if not versions:
        return {}
    return versions[-1].get("package", {}) or {}


def _extract_m1_mappings(m1_pkg: dict, m2_pkg: Optional[dict]) -> list:
    """Reconstruct m1_class_mappings from M1 class.parent_class_name and M2 level enums.

    Because mappings aren't stored explicitly after save-m2, we re-derive them
    from the persisted state: each M1 class's parent_class_name + any `level`
    default_value on the inherited attribute.
    """
    if not m2_pkg or not m1_pkg:
        return []

    m2_classes = {c.get("name", ""): c for c in (m2_pkg.get("classes") or [])}
    enums_by_id = {e.get("id"): e for e in (m2_pkg.get("enumerations") or [])}

    mappings = []
    for cls in (m1_pkg.get("classes") or []):
        parent = cls.get("parent_class_name")
        if not parent or parent not in m2_classes:
            continue

        level = None
        level_enum_id = None
        # Recover level assignment from M1's `level` attribute's default_value
        for a in (cls.get("attributes") or []):
            if a.get("name") == "level" and a.get("is_inherited"):
                level_enum_id = a.get("enum_ref")
                if a.get("default_value"):
                    level = a["default_value"]
                break

        # Fallback: if M2 has a level enum but M1 has no default, mark whole_tree
        m2_has_level = False
        for a in (m2_classes[parent].get("attributes") or []):
            if a.get("name") == "level" and a.get("enum_ref") in enums_by_id:
                m2_has_level = True
                level_enum_id = a.get("enum_ref")
                break
        if m2_has_level and not level:
            level = "whole_tree"

        entry = {"m1_class_name": cls.get("name", ""), "m2_parent_name": parent}
        if level:
            entry["level"] = level
        if level_enum_id:
            entry["m2_level_enum_id"] = level_enum_id
        mappings.append(entry)

    return mappings


# ==============================================================
# Word report (.docx)
# ==============================================================

def _build_word(m1_model, m2_model, m1_pkg, m2_pkg, mappings, svg_by_name: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Set default font ---
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # --- Cover page ---
    title = doc.add_heading("数据模型审查报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m1_label = m1_model.get("label") or m1_model.get("name") or "M1"
    m2_label = (m2_model.get("label") or m2_model.get("name")) if m2_model else "(未生成)"
    m1_classes = m1_pkg.get("classes") or []
    m2_classes = (m2_pkg.get("classes") or []) if m2_pkg else []

    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(stats_p, f"\nM1 模型:  {m1_label}\n", bold=True, size=14)
    _add_run(stats_p, f"M2 元模型: {m2_label}\n", bold=True, size=14)
    _add_run(stats_p, f"\n生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n", size=11)
    _add_run(
        stats_p,
        f"\n共识别 {len(m2_classes)} 个 M2 业务主题 · {len(m1_classes)} 个 M1 领域类\n",
        size=12,
    )

    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading("一、执行摘要", level=1)
    total_attrs = sum(
        len([a for a in (c.get("attributes") or []) if not a.get("is_inherited")])
        for c in m1_classes
    )
    n_hierarchy = 0
    if m2_pkg:
        for c in m2_classes:
            if any(
                a.get("name") == "level" and a.get("enum_ref")
                for a in (c.get("attributes") or [])
            ):
                n_hierarchy += 1
    total_assocs = len((m1_pkg.get("associations") or []))

    p = doc.add_paragraph()
    _add_run(p, "本次模型构建成果:\n", bold=True)
    doc.add_paragraph(f"· M2 元模型共识别 {len(m2_classes)} 个业务主题基类，"
                      f"其中 {n_hierarchy} 个主题具有多层级结构", style="List Bullet")
    doc.add_paragraph(f"· M1 领域模型覆盖 {len(m1_classes)} 个具体业务类，"
                      f"累计 {total_attrs} 个自有属性", style="List Bullet")
    doc.add_paragraph(f"· M1 之间识别出 {total_assocs} 条关联关系", style="List Bullet")

    doc.add_paragraph(
        "建议优先核对以下事项:", style="Intense Quote"
    ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    for point in [
        "业务主题的划分是否符合你们的分析维度？",
        "层级结构 (如 设施→功能分组→设备→部件) 是否合理？",
        "每个 M1 类归属到的 M2 主题和层级是否正确？",
        "关键属性是否完整？有无冗余或遗漏？",
        "空属性或孤立的 M2 基类是否需要合并或删除？",
    ]:
        doc.add_paragraph(f"  ◦ {point}", style="List Bullet")

    doc.add_page_break()

    # --- Section 2: M2 Overview ---
    doc.add_heading("二、M2 元模型一览", level=1)

    if not m2_pkg or not m2_classes:
        doc.add_paragraph("（尚未生成 M2 元模型）")
    else:
        doc.add_paragraph(
            f"本次共抽象出 {len(m2_classes)} 个 M2 业务主题基类。每个基类代表一个独立的"
            "业务观测维度，所有归属其下的 M1 类共享其定义的契约。"
        )

        # Embed overview SVG if available
        _insert_svg_if_available(doc, svg_by_name, 0, caption="图 2-1: M2 元模型总览")

        # M2 summary table
        doc.add_heading("2.1 M2 基类清单", level=2)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, text in enumerate(["业务主题", "中文标签", "M1 子类数", "是否带层级", "说明"]):
            hdr[i].text = text
            _set_cell_bold(hdr[i])

        child_count = _count_by_m2(mappings)
        for c in m2_classes:
            cname = c.get("name", "")
            row = t.add_row().cells
            row[0].text = cname
            row[1].text = c.get("label", "") or ""
            row[2].text = str(child_count.get(cname, 0))
            has_h = _m2_has_hierarchy(c, m2_pkg)
            row[3].text = "✓ 是" if has_h else "✗ 否"
            row[4].text = (c.get("description") or "")[:60]

        # 2.2+ per-M2 detail sections
        doc.add_page_break()
        for idx, c in enumerate(m2_classes):
            cname = c.get("name", "")
            clabel = c.get("label", "") or cname
            doc.add_heading(f"2.{idx + 2} {clabel} ({cname})", level=2)

            if c.get("description"):
                doc.add_paragraph(c["description"])

            # Levels
            levels = _get_m2_levels(c, m2_pkg)
            if levels:
                p = doc.add_paragraph()
                _add_run(p, "层级结构: ", bold=True)
                _add_run(p, " → ".join(levels))

            # Embed per-theme SVG if available (index i+1 for M2 themes)
            _insert_svg_if_available(
                doc, svg_by_name, idx + 1,
                caption=f"图 2-{idx + 2}: {clabel} · 主题归属",
            )

            # M2 attributes table
            attrs = c.get("attributes") or []
            if attrs:
                p = doc.add_paragraph()
                _add_run(p, "共性属性:", bold=True)
                at = doc.add_table(rows=1, cols=5)
                at.style = "Light Grid Accent 1"
                for i, text in enumerate(["属性名", "中文", "类型", "多重性", "说明"]):
                    cell = at.rows[0].cells[i]
                    cell.text = text
                    _set_cell_bold(cell)
                for a in attrs:
                    r = at.add_row().cells
                    r[0].text = a.get("name", "")
                    r[1].text = a.get("label", "") or ""
                    r[2].text = _fmt_type(a)
                    r[3].text = _fmt_mult(a.get("multiplicity"))
                    r[4].text = (a.get("description") or "")[:40]

            # M1 children grouped by level
            children = [m for m in mappings if m.get("m2_parent_name") == cname]
            if children:
                p = doc.add_paragraph()
                _add_run(p, f"包含的 M1 子类 ({len(children)} 个):", bold=True)
                if levels:
                    # ASCII tree by level
                    by_level = {}
                    orphans = []
                    for m in children:
                        lvl = m.get("level")
                        if lvl and lvl != "whole_tree" and lvl in levels:
                            by_level.setdefault(lvl, []).append(m)
                        else:
                            orphans.append(m)
                    for i, lvl in enumerate(levels):
                        group = by_level.get(lvl, [])
                        doc.add_paragraph(
                            f"├─ {lvl} 层 ({len(group)} 类)",
                            style="Normal",
                        )
                        for m in group:
                            m1_name = m.get("m1_class_name", "")
                            m1_cls = _find_m1(m1_pkg, m1_name)
                            m1_label = m1_cls.get("label", "") if m1_cls else ""
                            connector = "│  └─" if i < len(levels) - 1 else "   └─"
                            doc.add_paragraph(
                                f"{connector} {m1_label or m1_name} ({m1_name})",
                                style="Normal",
                            )
                    if orphans:
                        doc.add_paragraph(
                            f"└─ 全树模板 (whole_tree, {len(orphans)} 类)",
                            style="Normal",
                        )
                        for m in orphans:
                            m1_name = m.get("m1_class_name", "")
                            m1_cls = _find_m1(m1_pkg, m1_name)
                            m1_label = m1_cls.get("label", "") if m1_cls else ""
                            doc.add_paragraph(
                                f"   └─ {m1_label or m1_name} ({m1_name})",
                                style="Normal",
                            )
                else:
                    # Flat list
                    for m in children:
                        m1_name = m.get("m1_class_name", "")
                        m1_cls = _find_m1(m1_pkg, m1_name)
                        m1_label = m1_cls.get("label", "") if m1_cls else ""
                        doc.add_paragraph(
                            f"  ◦ {m1_label or m1_name} ({m1_name})",
                            style="List Bullet",
                        )

            doc.add_page_break()

    # --- Section 3: M1 领域模型 ---
    doc.add_heading("三、M1 领域模型清单", level=1)
    doc.add_paragraph(
        f"本节按 M2 主题分组，详细列出 {len(m1_classes)} 个 M1 领域类的属性与关联。"
    )

    # Group by M2 parent
    groups_by_m2 = {}
    for cls in m1_classes:
        parent = cls.get("parent_class_name") or "__unmapped__"
        groups_by_m2.setdefault(parent, []).append(cls)

    sec_idx = 1
    for m2_name, classes_in_group in groups_by_m2.items():
        if m2_name == "__unmapped__":
            doc.add_heading(f"3.{sec_idx} (未归属到任何 M2 主题)", level=2)
        else:
            m2_cls = next((x for x in m2_classes if x.get("name") == m2_name), None)
            m2_label = (m2_cls.get("label") if m2_cls else m2_name) or m2_name
            doc.add_heading(f"3.{sec_idx} {m2_label} 主题下的 M1 类", level=2)
        sec_idx += 1

        for cls in classes_in_group:
            cname = cls.get("name", "")
            clabel = cls.get("label", "") or cname
            p = doc.add_paragraph()
            _add_run(p, f"● {clabel} ", bold=True)
            _add_run(p, f"({cname})", size=9)

            meta = []
            if cls.get("parent_class_name"):
                meta.append(f"继承自: {cls['parent_class_name']}")
            # Level
            for a in (cls.get("attributes") or []):
                if a.get("name") == "level" and a.get("default_value"):
                    meta.append(f"默认层级: {a['default_value']}")
                    break
            if cls.get("description"):
                meta.append(f"描述: {cls['description'][:80]}")
            if meta:
                doc.add_paragraph("  " + " · ".join(meta))

            own_attrs = [a for a in (cls.get("attributes") or []) if not a.get("is_inherited")]
            inh_attrs = [a for a in (cls.get("attributes") or []) if a.get("is_inherited")]

            if own_attrs:
                t = doc.add_table(rows=1, cols=5)
                t.style = "Light Grid"
                for i, text in enumerate(["属性名", "中文", "类型", "单位", "多重性"]):
                    cell = t.rows[0].cells[i]
                    cell.text = text
                    _set_cell_bold(cell)
                for a in own_attrs:
                    r = t.add_row().cells
                    r[0].text = a.get("name", "")
                    r[1].text = a.get("label", "") or ""
                    r[2].text = _fmt_type(a)
                    r[3].text = a.get("unit", "") or ""
                    r[4].text = _fmt_mult(a.get("multiplicity"))

            if inh_attrs:
                p = doc.add_paragraph()
                _add_run(p, f"  (还继承 {len(inh_attrs)} 个属性，包括: ", italic=True, size=9)
                _add_run(p, ", ".join(a.get("name", "") for a in inh_attrs[:5]), italic=True, size=9)
                if len(inh_attrs) > 5:
                    _add_run(p, f" ... 共 {len(inh_attrs)} 个)", italic=True, size=9)
                else:
                    _add_run(p, ")", italic=True, size=9)

            doc.add_paragraph()  # breather

    # --- Section 4: Enumerations ---
    m1_enums = (m1_pkg.get("enumerations") or [])
    m2_enums = (m2_pkg.get("enumerations") or []) if m2_pkg else []
    if m1_enums or m2_enums:
        doc.add_page_break()
        doc.add_heading("四、枚举字典", level=1)

        all_enums = [(e, "M1") for e in m1_enums] + [(e, "M2") for e in m2_enums]
        for enum, layer in all_enums:
            ename = enum.get("name", "")
            elabel = enum.get("label", "") or ename
            p = doc.add_paragraph()
            _add_run(p, f"[{layer}] {elabel} ", bold=True)
            _add_run(p, f"({ename})", size=9)
            literals = enum.get("literals") or []
            if literals:
                lit_line = "  字面值: " + "、".join(
                    f"{l.get('label') or l.get('name')}" for l in literals
                )
                doc.add_paragraph(lit_line)

    # --- Appendix ---
    doc.add_page_break()
    doc.add_heading("附录 A: 配套审查意见表使用说明", level=1)
    for line in [
        "本报告配套的「审查意见表.xlsx」是填写反馈的主要工具。",
        "",
        "审查表共含 5 个工作表:",
        "  · 总体意见 — 填写综合评价",
        "  · M2 基类审查 — 每个 M2 主题一行, 下拉选择【通过/需修改/拒绝】",
        "  · 层级结构审查 — 对有层级的 M2 主题, 核对层级序列",
        "  · M1 类审查 — 每个 M1 类一行, 核对归属和层级分配",
        "  · 关联与枚举审查 — 检查关联关系和枚举定义",
        "",
        "每行右侧有固定的【决策】(下拉) 和【审查意见】(自由文本) 列。",
        "填完回传后, 系统可以自动导入你的反馈并标注到对应实体上。",
        "",
        "⚠ 请勿修改表格中 id 列的内容 — 系统依赖该列回读反馈。",
    ]:
        doc.add_paragraph(line)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==============================================================
# Excel workbook (.xlsx)
# ==============================================================

def _build_excel(m1_model, m2_model, m1_pkg, m2_pkg, mappings) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4F81BD")
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    wrap = Alignment(wrap_text=True, vertical="top")

    m1_label = m1_model.get("label") or m1_model.get("name") or "M1"

    # =========== Sheet 0: 总体意见 ===========
    ws0 = wb.create_sheet("总体意见")
    ws0.column_dimensions["A"].width = 20
    ws0.column_dimensions["B"].width = 60
    ws0["A1"] = "字段"
    ws0["B1"] = "内容"
    ws0["A1"].font = header_font
    ws0["B1"].font = header_font
    ws0["A1"].fill = header_fill
    ws0["B1"].fill = header_fill

    rows = [
        ("审查项目", f"{m1_label} 数据模型"),
        ("审查人姓名", ""),
        ("审查日期", ""),
        ("业务部门", ""),
        ("", ""),
        ("总体评分 (1-5分)", ""),
        ("整体业务主题划分是否合理？", ""),
        ("层级结构设计是否合理？", ""),
        ("属性完整度评价", ""),
        ("", ""),
        ("综合意见 (可长篇)", ""),
        ("后续改进建议", ""),
    ]
    for i, (k, v) in enumerate(rows, start=2):
        ws0.cell(row=i, column=1, value=k).font = Font(bold=bool(k and not k.startswith(" ")))
        ws0.cell(row=i, column=2, value=v)
        ws0.cell(row=i, column=2).alignment = wrap

    # Highlight editable cells
    editable_rows = [3, 4, 5, 7, 8, 9, 10, 12, 13]
    light_yellow = PatternFill("solid", fgColor="FFF2CC")
    for r in editable_rows:
        ws0.cell(row=r, column=2).fill = light_yellow

    # =========== Sheet 1: M2 基类审查 ===========
    ws1 = wb.create_sheet("M2基类审查")
    headers1 = ["id", "M2 主题 (英文)", "中文标签", "M1 子类数", "是否带层级",
                "描述", "决策", "审查意见", "修改建议"]
    for i, h in enumerate(headers1, start=1):
        c = ws1.cell(row=1, column=i, value=h)
        c.font = header_font
        c.fill = header_fill
    widths1 = [30, 25, 20, 10, 10, 40, 12, 40, 40]
    for i, w in enumerate(widths1, start=1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    m2_classes = (m2_pkg.get("classes") or []) if m2_pkg else []
    child_count = _count_by_m2(mappings)
    for r, c in enumerate(m2_classes, start=2):
        cname = c.get("name", "")
        ws1.cell(row=r, column=1, value=f"m2:{cname}")
        ws1.cell(row=r, column=2, value=cname)
        ws1.cell(row=r, column=3, value=c.get("label", ""))
        ws1.cell(row=r, column=4, value=child_count.get(cname, 0))
        ws1.cell(row=r, column=5, value="是" if _m2_has_hierarchy(c, m2_pkg) else "否")
        ws1.cell(row=r, column=6, value=(c.get("description") or "")[:100])
        # Highlight decision/comment cells
        for col in (7, 8, 9):
            ws1.cell(row=r, column=col).fill = light_yellow
            ws1.cell(row=r, column=col).alignment = wrap
    # Add data validation for decision column (col 7)
    dv1 = DataValidation(type="list", formula1='"通过,需修改,拒绝,待议"', allow_blank=True)
    ws1.add_data_validation(dv1)
    if len(m2_classes) > 0:
        dv1.add(f"G2:G{len(m2_classes) + 1}")

    # =========== Sheet 2: 层级结构审查 ===========
    ws2 = wb.create_sheet("层级结构审查")
    headers2 = ["id", "M2 主题", "中文", "层级序列", "层级数", "M1 分配统计",
                "决策", "层级是否合理", "调整建议"]
    for i, h in enumerate(headers2, start=1):
        c = ws2.cell(row=1, column=i, value=h)
        c.font = header_font
        c.fill = header_fill
    widths2 = [30, 25, 20, 40, 8, 25, 12, 40, 40]
    for i, w in enumerate(widths2, start=1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    hier_rows = 0
    for c in m2_classes:
        levels = _get_m2_levels(c, m2_pkg)
        if not levels:
            continue
        hier_rows += 1
        r = hier_rows + 1
        cname = c.get("name", "")
        ws2.cell(row=r, column=1, value=f"m2_level:{cname}")
        ws2.cell(row=r, column=2, value=cname)
        ws2.cell(row=r, column=3, value=c.get("label", ""))
        ws2.cell(row=r, column=4, value=" → ".join(levels))
        ws2.cell(row=r, column=5, value=len(levels))
        # Count M1 per level
        per_level = {}
        for m in mappings:
            if m.get("m2_parent_name") == cname:
                lvl = m.get("level") or "未分配"
                per_level[lvl] = per_level.get(lvl, 0) + 1
        dist = "、".join(f"{k}({v})" for k, v in per_level.items())
        ws2.cell(row=r, column=6, value=dist)
        for col in (7, 8, 9):
            ws2.cell(row=r, column=col).fill = light_yellow
            ws2.cell(row=r, column=col).alignment = wrap

    dv2 = DataValidation(type="list", formula1='"通过,需修改,拒绝,待议"', allow_blank=True)
    ws2.add_data_validation(dv2)
    if hier_rows > 0:
        dv2.add(f"G2:G{hier_rows + 1}")

    # =========== Sheet 3: M1类审查 ===========
    ws3 = wb.create_sheet("M1类审查")
    headers3 = ["id", "M1 类名", "中文标签", "M2 父类", "层级", "自有属性数", "关联数",
                "描述", "决策", "归属是否合理", "修改建议"]
    for i, h in enumerate(headers3, start=1):
        c = ws3.cell(row=1, column=i, value=h)
        c.font = header_font
        c.fill = header_fill
    widths3 = [30, 28, 22, 22, 15, 10, 8, 30, 12, 30, 35]
    for i, w in enumerate(widths3, start=1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    m1_classes = m1_pkg.get("classes") or []
    # count assoc per class
    assoc_count_by_cls = {}
    for a in (m1_pkg.get("associations") or []):
        src = (a.get("source") or {}).get("class_name")
        tgt = (a.get("target") or {}).get("class_name")
        if src:
            assoc_count_by_cls[src] = assoc_count_by_cls.get(src, 0) + 1
        if tgt and tgt != src:
            assoc_count_by_cls[tgt] = assoc_count_by_cls.get(tgt, 0) + 1

    mapping_by_m1 = {m.get("m1_class_name"): m for m in mappings}

    for r, cls in enumerate(m1_classes, start=2):
        cname = cls.get("name", "")
        mapping = mapping_by_m1.get(cname)
        ws3.cell(row=r, column=1, value=f"m1:{cname}")
        ws3.cell(row=r, column=2, value=cname)
        ws3.cell(row=r, column=3, value=cls.get("label", ""))
        ws3.cell(row=r, column=4, value=(mapping or {}).get("m2_parent_name", cls.get("parent_class_name") or ""))
        ws3.cell(row=r, column=5, value=(mapping or {}).get("level", "") or "")
        ws3.cell(row=r, column=6,
                 value=len([a for a in (cls.get("attributes") or []) if not a.get("is_inherited")]))
        ws3.cell(row=r, column=7, value=assoc_count_by_cls.get(cname, 0))
        ws3.cell(row=r, column=8, value=(cls.get("description") or "")[:80])
        for col in (9, 10, 11):
            ws3.cell(row=r, column=col).fill = light_yellow
            ws3.cell(row=r, column=col).alignment = wrap

    dv3 = DataValidation(type="list", formula1='"通过,需修改,拒绝,待议"', allow_blank=True)
    ws3.add_data_validation(dv3)
    if len(m1_classes) > 0:
        dv3.add(f"I2:I{len(m1_classes) + 1}")

    # =========== Sheet 4: 关联与枚举审查 ===========
    ws4 = wb.create_sheet("关联与枚举审查")
    headers4 = ["id", "类型", "名称", "详情", "决策", "审查意见"]
    for i, h in enumerate(headers4, start=1):
        c = ws4.cell(row=1, column=i, value=h)
        c.font = header_font
        c.fill = header_fill
    widths4 = [35, 10, 30, 70, 12, 40]
    for i, w in enumerate(widths4, start=1):
        ws4.column_dimensions[get_column_letter(i)].width = w

    row_num = 2
    for a in (m1_pkg.get("associations") or []):
        aid = a.get("id", "") or a.get("name", "")
        ws4.cell(row=row_num, column=1, value=f"assoc_m1:{aid}")
        ws4.cell(row=row_num, column=2, value="M1关联")
        ws4.cell(row=row_num, column=3, value=a.get("label", "") or a.get("name", ""))
        src = (a.get("source") or {}).get("class_name", "?")
        tgt = (a.get("target") or {}).get("class_name", "?")
        atype = a.get("association_type", "association")
        detail = f"{src} →({atype})→ {tgt}"
        ws4.cell(row=row_num, column=4, value=detail)
        for col in (5, 6):
            ws4.cell(row=row_num, column=col).fill = light_yellow
            ws4.cell(row=row_num, column=col).alignment = wrap
        row_num += 1

    # Enums (M1 + M2)
    for enum_list, prefix, tag in [
        ((m1_pkg.get("enumerations") or []), "enum_m1", "M1枚举"),
        (((m2_pkg or {}).get("enumerations") or []), "enum_m2", "M2枚举"),
    ]:
        for e in enum_list:
            eid = e.get("id", "") or e.get("name", "")
            ws4.cell(row=row_num, column=1, value=f"{prefix}:{eid}")
            ws4.cell(row=row_num, column=2, value=tag)
            ws4.cell(row=row_num, column=3, value=e.get("label", "") or e.get("name", ""))
            literals = e.get("literals") or []
            lit_str = "、".join(l.get("label") or l.get("name", "") for l in literals[:8])
            if len(literals) > 8:
                lit_str += f"... (共{len(literals)}项)"
            ws4.cell(row=row_num, column=4, value=lit_str)
            for col in (5, 6):
                ws4.cell(row=row_num, column=col).fill = light_yellow
                ws4.cell(row=row_num, column=col).alignment = wrap
            row_num += 1

    dv4 = DataValidation(type="list", formula1='"通过,需修改,拒绝,待议"', allow_blank=True)
    ws4.add_data_validation(dv4)
    if row_num > 2:
        dv4.add(f"E2:E{row_num - 1}")

    # Freeze headers and apply borders to header row on all sheets
    for ws in (ws0, ws1, ws2, ws3, ws4):
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.border = thin_border

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ==============================================================
# Source JSON + README
# ==============================================================

def _build_source_json(m1_model, m2_model, mappings) -> str:
    data = {
        "exported_at": datetime.now().isoformat(),
        "m1_model": m1_model,
        "m2_model": m2_model,
        "m1_class_mappings": mappings,
    }
    # model.model_dump() preserves datetime objects; json needs a fallback
    return json.dumps(data, ensure_ascii=False, indent=2, default=_json_default)


def _json_default(obj):
    from datetime import datetime as _dt, date as _date
    if isinstance(obj, (_dt, _date)):
        return obj.isoformat()
    return str(obj)


def _build_readme(m1_model, m2_model) -> str:
    m1_label = m1_model.get("label") or m1_model.get("name") or "M1"
    m2_label = (m2_model.get("label") or m2_model.get("name")) if m2_model else "(未生成)"
    return f"""# 审查包使用说明

## 📋 本包内容

| 文件 | 作用 | 推荐打开工具 |
|------|------|--------------|
| `审查报告.docx` | 叙述式审查报告 (含图) | Microsoft Word |
| `审查意见表.xlsx` | 结构化反馈表格 | Microsoft Excel |
| `审查图集.pdf` | 关系图静态版 (多页 PDF) | **任何 PDF 阅读器** (Adobe / 浏览器 / Word) |
| `审查图集.drawio` | 可编辑关系图 (多页) | draw.io / app.diagrams.net (可选) |
| `图片快照/*.svg` | 关系图单页高清矢量 | 浏览器 / Word (365+) |
| `附件_源数据.json` | 原始结构化数据 | 文本编辑器 / 开发 |

## ✅ 审查流程 (推荐)

### 第一步: 宏观浏览 (30 分钟)
1. 打开 **审查报告.docx**
2. 先看「执行摘要」了解总体情况
3. 看「二、M2 元模型一览」核对业务主题划分是否合理
4. 滚动到各 M2 主题的详情, 核对层级结构

### 第二步: 视觉核对 (20 分钟)
**推荐路径 (最省事)**: 打开 **审查图集.pdf** — 任何 PDF 阅读器都能看, 多页滑动浏览即可。

**可编辑路径 (需要 draw.io)**: 打开 **审查图集.drawio**, 用 https://app.diagrams.net/ 在线打开或本地装 draw.io。可以在图上画箭头、写批注、拖动节点调整。

两种方式内容一样, 按喜好选:
1. 先看 M2 总览页, 确认全貌
2. 翻到你负责的业务主题页, 逐个核对 M1 归属和层级分配

### 第三步: 填写反馈 (主要工作, 1-2 小时)
1. 打开 **审查意见表.xlsx**
2. 先在「总体意见」表填写综合评价
3. 在「M2 基类审查」逐行审核, 下拉选择决策 (通过/需修改/拒绝/待议), 写明意见
4. 同样处理「层级结构审查」「M1 类审查」「关联与枚举审查」
5. 保存并回传

## ⚠️ 填写须知

- **不要修改 id 列** — 系统依赖它解析你的反馈
- **决策列有下拉框** — 尽量使用下拉值, 不要自己打别的字
- **黄色单元格是需要你填写的**
- 对每一条「需修改」或「拒绝」的, 请在「修改建议」列写明具体怎么改

## 📦 模型信息

- **M1 模型**: {m1_label}
- **M2 元模型**: {m2_label}
- **生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## ❓ 技术问题

- 不会用 drawio?  不用管它, 直接看 **审查图集.pdf** 即可
- SVG 打不开?  用最新版 Chrome / Edge / Word 365
- Drawio 装不了?  访问 https://app.diagrams.net/ (免费在线版, 不需安装)
- JSON 不用管它, 这是给开发人员的备份

---

感谢审查! 你的反馈会帮助我们把数据模型做得更好。
"""


# ==============================================================
# Utility helpers
# ==============================================================

def _count_by_m2(mappings: list) -> dict:
    out = {}
    for m in (mappings or []):
        p = m.get("m2_parent_name")
        if p:
            out[p] = out.get(p, 0) + 1
    return out


def _m2_has_hierarchy(m2_cls: dict, m2_pkg: dict) -> bool:
    enum_ids = {e.get("id") for e in (m2_pkg.get("enumerations") or [])}
    for a in (m2_cls.get("attributes") or []):
        if a.get("name") == "level" and a.get("enum_ref") in enum_ids:
            return True
    return False


def _get_m2_levels(m2_cls: dict, m2_pkg: dict) -> list:
    enum_id = None
    for a in (m2_cls.get("attributes") or []):
        if a.get("name") == "level":
            enum_id = a.get("enum_ref")
            break
    if not enum_id:
        return []
    for e in (m2_pkg.get("enumerations") or []):
        if e.get("id") == enum_id:
            return [l.get("label") or l.get("name") for l in (e.get("literals") or [])]
    return []


def _find_m1(m1_pkg: dict, name: str):
    for c in (m1_pkg.get("classes") or []):
        if c.get("name") == name:
            return c
    return None


def _fmt_type(a: dict) -> str:
    t = a.get("data_type") or "String"
    unit = a.get("unit")
    return f"{t} ({unit})" if unit else t


def _fmt_mult(m):
    if not m:
        return "1..1"
    lower = m.get("lower", 1) if isinstance(m, dict) else 1
    upper = m.get("upper", 1) if isinstance(m, dict) else 1
    u = "*" if upper == -1 else str(upper)
    return f"{lower}..{u}"


# ==============================================================
# Word cell/run helpers
# ==============================================================

def _add_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    from docx.shared import Pt, RGBColor
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _set_cell_bold(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _insert_svg_if_available(doc, svg_by_name: dict, index: int, caption: str = ""):
    """Rasterize the SVG at the given index to PNG and embed it directly in the docx.

    python-docx has no native SVG support, so we convert SVG -> PNG on the fly using
    svglib + reportlab.renderPM (with the same CJK font forcing used for the PDF, to
    keep Chinese characters from becoming tofu). The embedded PNG is sized to ~6.5
    inches wide so it fits a standard A4/Letter page with margins. If rasterization
    fails (missing libs or unparseable SVG), we fall back to a text reference so the
    user can still find the image in the 图片快照/ folder.
    """
    import io as _io
    from docx.shared import Pt, Inches
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        _CENTER = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        _CENTER = 1  # numeric fallback

    if not svg_by_name:
        return
    ordered_names = list(svg_by_name.keys())
    if index >= len(ordered_names):
        return

    fname = ordered_names[index]
    svg_str = svg_by_name[fname]

    # Try to rasterize to PNG
    png_bytes = None
    try:
        from backend.services.diagram_exporter import svg_to_png_bytes
        png_bytes = svg_to_png_bytes(svg_str, dpi=144)
    except Exception:
        png_bytes = None

    if png_bytes:
        # Embed as actual image (centered, ~6.5 inches wide)
        p = doc.add_paragraph()
        p.alignment = _CENTER
        run = p.add_run()
        try:
            run.add_picture(_io.BytesIO(png_bytes), width=Inches(6.3))
        except Exception:
            # If add_picture fails for any reason, fall through to text marker
            p.clear()
            run = p.add_run(f"[见图: 图片快照/{fname}]")
            run.italic = True
            run.font.size = Pt(10)
        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = _CENTER
            _add_run(cap_p, caption, bold=True, size=10)
        return

    # Fallback: text reference (reader must open the SVG in the 图片快照/ folder)
    p = doc.add_paragraph()
    p.alignment = _CENTER
    run = p.add_run(f"[见图: 图片快照/{fname}]  (自动图像嵌入失败, 请打开图片快照文件夹查看)")
    run.italic = True
    run.font.size = Pt(10)
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = _CENTER
        _add_run(cap_p, caption, bold=True, size=10)

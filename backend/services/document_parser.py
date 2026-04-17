"""Extract text content from uploaded documents (TXT, MD, PDF, DOCX, Excel)."""
from __future__ import annotations

from pathlib import Path


class DocumentParser:

    def parse(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        handlers = {
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".xlsx": self._parse_excel,
            ".xls": self._parse_excel,
            ".csv": self._parse_csv,
        }
        handler = handlers.get(suffix)
        if handler is None:
            raise ValueError(f"Unsupported file format: {suffix}")
        return handler(file_path)

    def _parse_text(self, path: Path) -> str:
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                return path.read_text(encoding=enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return path.read_text(encoding="utf-8", errors="replace")

    def _parse_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError("pdfplumber is required for PDF parsing. pip install pdfplumber")

        pages_text = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                # Extract tables as markdown
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        text += "\n\n"
                        text += self._table_to_markdown(table)
                if text.strip():
                    pages_text.append(f"--- Page {i+1} ---\n{text}")
        return "\n\n".join(pages_text)

    def _parse_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx is required for DOCX parsing. pip install python-docx")

        doc = Document(str(path))
        parts = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                for para in doc.paragraphs:
                    if para._element is element:
                        style = para.style.name if para.style else ""
                        prefix = ""
                        if "Heading 1" in style:
                            prefix = "# "
                        elif "Heading 2" in style:
                            prefix = "## "
                        elif "Heading 3" in style:
                            prefix = "### "
                        if para.text.strip():
                            parts.append(f"{prefix}{para.text}")
                        break

            elif tag == "tbl":
                # Table
                for table in doc.tables:
                    if table._element is element:
                        rows = []
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            rows.append(cells)
                        if rows:
                            parts.append(self._table_to_markdown(rows))
                        break

        return "\n\n".join(parts)

    def _parse_excel(self, path: Path) -> str:
        try:
            import openpyxl
        except ImportError:
            raise RuntimeError("openpyxl is required for Excel parsing. pip install openpyxl")

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(cells)
            if rows:
                parts.append(f"## Sheet: {sheet_name}\n")
                parts.append(self._table_to_markdown(rows))
        wb.close()
        return "\n\n".join(parts)

    def _parse_csv(self, path: Path) -> str:
        import csv
        text = self._parse_text(path)
        rows = []
        for row in csv.reader(text.splitlines()):
            rows.append(row)
        if rows:
            return self._table_to_markdown(rows)
        return text

    def summarize_for_modeling(self, full_text: str, filename: str, max_chars: int = 3000) -> str:
        """
        Create a modeling-oriented summary of a document.
        For tables: keep headers + a few sample rows + column stats.
        For text: keep first portion + structure hints.
        This is what gets sent to the AI — NOT the full text.
        """
        if len(full_text) <= max_chars:
            return full_text

        lines = full_text.split("\n")

        # Detect if this is primarily tabular (markdown tables)
        table_lines = [l for l in lines if l.strip().startswith("|")]
        if len(table_lines) > len(lines) * 0.3:
            # Tabular document — extract schema per table
            return self._summarize_tables(lines, filename, max_chars)
        else:
            # Text document — keep beginning and structure
            return self._summarize_text(full_text, filename, max_chars)

    def _summarize_tables(self, lines: list[str], filename: str, max_chars: int) -> str:
        """Extract table headers + sample rows + stats from table-heavy documents."""
        parts = [f"[文档: {filename}]"]
        current_header = ""
        current_table_rows = []
        tables_found = 0

        for line in lines:
            stripped = line.strip()
            # Section headers
            if stripped.startswith("#"):
                if current_table_rows:
                    parts.append(self._compress_table(current_header, current_table_rows))
                    tables_found += 1
                    current_table_rows = []
                current_header = stripped
            elif stripped.startswith("|") and "---" not in stripped:
                current_table_rows.append(stripped)
            # Skip separator rows (|---|---|)

        # Last table
        if current_table_rows:
            parts.append(self._compress_table(current_header, current_table_rows))
            tables_found += 1

        result = "\n\n".join(parts)
        if len(result) > max_chars:
            result = result[:max_chars] + f"\n\n[... 截断, 原文 {len(''.join(lines)):,} 字符, {tables_found} 个表]"
        return result

    @staticmethod
    def _compress_table(header: str, rows: list[str]) -> str:
        """Keep header row + up to 5 sample rows + row count."""
        total = len(rows)
        if total == 0:
            return ""
        parts = []
        if header:
            parts.append(header)

        # First row is the table header
        parts.append(rows[0])
        # Separator
        cols = rows[0].count("|") - 1
        parts.append("| " + " | ".join(["---"] * max(cols, 1)) + " |")
        # Sample rows: first 3 + last 2 (if enough rows)
        data_rows = rows[1:]
        if len(data_rows) <= 5:
            parts.extend(data_rows)
        else:
            parts.extend(data_rows[:3])
            parts.append(f"| ... 省略 {len(data_rows) - 5} 行 ... |")
            parts.extend(data_rows[-2:])

        parts.append(f"[共 {total} 行数据]")
        return "\n".join(parts)

    @staticmethod
    def _summarize_text(full_text: str, filename: str, max_chars: int) -> str:
        """Keep beginning of text documents with a truncation note."""
        keep = max_chars - 200
        return (
            f"[文档: {filename}]\n\n"
            + full_text[:keep]
            + f"\n\n[... 截断, 原文共 {len(full_text):,} 字符]"
        )

    @staticmethod
    def _table_to_markdown(rows: list[list]) -> str:
        if not rows:
            return ""
        rows = [[str(c) if c is not None else "" for c in row] for row in rows]
        col_count = max(len(row) for row in rows)
        rows = [row + [""] * (col_count - len(row)) for row in rows]

        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)


parser = DocumentParser()
